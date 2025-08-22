import os
import io
import tempfile
import warnings
import docx
import pandas as pd
from dotenv import load_dotenv
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import pytesseract
from PIL import Image

# Suppress deprecation warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=UserWarning)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains.retrieval_qa.base import RetrievalQA
from langchain_groq import ChatGroq
from langchain.docstore.document import Document
from langchain_core.prompts import PromptTemplate
# from langchain.schema import Document
from langchain.chains.summarize import load_summarize_chain


load_dotenv()

def initialize_components():
    try:
        embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            cache_folder="./model_cache",  # Specify a local cache directory
            encode_kwargs={'normalize_embeddings': True}
        )
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        return embeddings, text_splitter
    except OSError as e:
        st.error(f"Error initializing components: {str(e)}")
        st.info("Please check your internet connection and try again.")
        return None, None
    
def extract_text_from_file(uploaded_file):
    text = ""
    file_ext = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_ext == 'pdf':
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_path = temp_file.name
            
            pdf_reader = PdfReader(temp_path)
            text = "\n".join([page.extract_text() or "" for page in pdf_reader.pages])
            os.unlink(temp_path)  # Clean up temp file
            
        elif file_ext == 'docx':
            doc = Document(io.BytesIO(uploaded_file.getvalue()))
            text = "\n".join([para.text for para in doc.paragraphs])
            
        elif file_ext in ['png', 'jpg', 'jpeg']:
            image = Image.open(uploaded_file)
            text = pytesseract.image_to_string(image)
            
        elif file_ext == 'txt':
            text = str(uploaded_file.getvalue(), "utf-8")
            
        elif file_ext in ['xlsx', 'xls', 'csv']:
            if file_ext == 'csv':
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            
            # Store dataframe in session state for direct access
            st.session_state.excel_data = df
            
            # Convert dataframe to text for RAG processing
            text = df.to_string(index=False)
            # Add column names and data types as metadata to improve RAG
            text += "\n\nTable Structure:\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n"
            text += f"Number of records: {len(df)}\n"
            for col in df.columns:
                text += f"Column '{col}' has data type: {df[col].dtype}\n"
                if df[col].dtype in ['object', 'string']:
                    sample_values = df[col].dropna().unique()[:5]
                    if len(sample_values) > 0:
                        text += f"Sample values for '{col}': {', '.join(str(x) for x in sample_values)}\n"
            
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None
    return text if text.strip() else None    

def process_document(uploaded_file, text_splitter, embeddings):
    text = extract_text_from_file(uploaded_file)
    if not text:
        st.error("No readable text found in the document")
        return None
    
    chunks = text_splitter.split_text(text)
    if not chunks:
        st.error("Document processing resulted in zero chunks")
        return None
        
    try:
        vector_store = FAISS.from_texts(chunks, embeddings)
        return vector_store
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        return None

def loader(file_path):
    if file_path.endswith(".csv"):
        df = pd.read_csv(file_path)
        text = df.to_string()
        return [Document(page_content=text)]   # 👈 returns a list of Documents

    elif file_path.endswith(".pdf"):
        pdf = PdfReader(file_path)
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
        return [Document(page_content=text)]

    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text)]

    else:  # txt or other formats
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text)]

def summarize_document(docs, llm):
    try:
        summary_chain = load_summarize_chain(llm, chain_type="map_reduce")
        return summary_chain.run(docs)
    except Exception as e:
        return f"Error generating summary: {str(e)}"

def get_summarization_prompt():
    return PromptTemplate.from_template("""
    Summarize the following text in clear, concise language...
    Text:
    {context}
    """)


def get_custom_prompt():
    return PromptTemplate(
        template="""Use the following pieces of context to answer the question at the end. 
        If you don't know the answer, just say that you don't know, don't try to make up an answer.
        
        Context: {context}
        
        Question: {question}
        
        Answer in detail and be specific. If relevant, include page numbers or sections from the document:""",
        input_variables=["context", "question"]
    )

def main():
    st.set_page_config(
        page_title="Advanced Document Analysis with Groq",
        page_icon="🧠",
        layout="wide"
    )
    
    components = initialize_components()
    if components is None:
        st.stop()
    
    embeddings, text_splitter = components
    
    # Check for Tesseract OCR quietly
    tesseract_available = True
    try:
        version = pytesseract.get_tesseract_version()
        if not version:
            tesseract_available = False
    except Exception:
        tesseract_available = False
    
    st.title("🧠 Advanced Document Analysis with Groq")
    st.write("Upload any document and get AI-powered insights")
    
    # Initialize session state
    if 'vector_store' not in st.session_state:
        st.session_state.vector_store = None
    if 'processed_file' not in st.session_state:
        st.session_state.processed_file = None
    if 'excel_data' not in st.session_state:
        st.session_state.excel_data = None
    
    # Sidebar configuration
    with st.sidebar:
        st.header("Configuration")
        groq_api_key = st.text_input(
            "Groq API Key",
            type="password",
            value=os.getenv("GROQ_API_KEY", "")
        )
        
        if groq_api_key:
            os.environ["GROQ_API_KEY"] = groq_api_key
        else:
            st.warning("Please enter your Groq API Key to generate answers")
        
        model_name = st.selectbox(
            "Select Groq Model",
            ["llama3-70b-8192", "llama3-8b-8192", "mixtral-8x7b-4096", "gemma-7b-it"],
            index=0
        )
        
        temperature = st.slider(
            "Creativity (Temperature)",
            min_value=0.0,
            max_value=1.0,
            value=0.3,
            step=0.1
        )
        
        st.markdown("---")
        st.markdown("### Supported Formats")
        st.markdown("- PDF, DOCX, TXT, PNG, JPG")
        st.markdown("---")
        st.markdown("### How to Use")
        st.markdown("1. Upload a document\n2. Ask questions\n3. Get detailed answers")
    
    # File upload section
    uploaded_file = st.file_uploader(
        "Choose a document",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg", "xlsx", "xls", "csv"],
        accept_multiple_files=False,
        key="file_uploader"
    )
    
    # Show Tesseract warning only when image file is uploaded
    if uploaded_file and not tesseract_available and uploaded_file.name.split('.')[-1].lower() in ['png', 'jpg', 'jpeg']:
        st.info("📌 Note: Tesseract OCR is not detected. Text extraction from images may not work properly. For best results with images, consider installing Tesseract OCR.", icon="ℹ️")
    
    # Process document if new file uploaded
    if uploaded_file and (st.session_state.processed_file != uploaded_file.name):
        with st.spinner("Processing document..."):
            vector_store = process_document(uploaded_file, text_splitter, embeddings)
            
            if vector_store:
                st.session_state.vector_store = vector_store
                st.session_state.processed_file = uploaded_file.name
                st.success(f"Document processed successfully! ({len(vector_store.index_to_docstore_id)} chunks)")
            else:
                st.error("Failed to process document. Please try a different file.")
    
    # Question and answer section
    if st.session_state.vector_store:
        st.subheader("Ask About the Document")
        
        # Excel data preview if available
        file_ext = ""
        if uploaded_file:
            file_ext = uploaded_file.name.split('.')[-1].lower()
        
        if file_ext in ['xlsx', 'xls', 'csv'] and st.session_state.excel_data is not None:
            with st.expander("Preview Data", expanded=True):
                st.dataframe(st.session_state.excel_data, use_container_width=True)
                
                # Quick filter for Excel data
                if not st.session_state.excel_data.empty:
                    st.subheader("Quick Filter")
                    
                    # Select column to filter
                    cols = st.session_state.excel_data.columns.tolist()
                    filter_col = st.selectbox("Select column to filter", cols)
                    
                    # Get unique values for the selected column
                    unique_values = st.session_state.excel_data[filter_col].dropna().unique().tolist()
                    if len(unique_values) > 0:
                        selected_value = st.selectbox(f"Select value from {filter_col}", unique_values)
                        
                        # Filter data
                        filtered_data = st.session_state.excel_data[st.session_state.excel_data[filter_col] == selected_value]
                        st.write(f"Filtered data for {filter_col} = {selected_value}:")
                        st.dataframe(filtered_data, use_container_width=True)
                        
                        # Quick questions based on filtered data
                        employee_suggestion = f"Show me details about {selected_value}" if filter_col == "Name" else ""
                        summary_suggestion = f"Summarize information about {selected_value}"
                        st.write("Quick questions:")
                        quick_q_col1, quick_q_col2 = st.columns(2)
                        
                        with quick_q_col1:
                            if st.button(employee_suggestion or "Show details"):
                                st.session_state.question = employee_suggestion or f"Show all details for {filter_col} = {selected_value}"
                        
                        with quick_q_col2:
                            if st.button(summary_suggestion):
                                st.session_state.question = summary_suggestion
                        
        # Initialize question state if not exist
        if 'question' not in st.session_state:
            st.session_state.question = ""
            
        question = st.text_area(
            "Enter your question",
            value=st.session_state.question,
            placeholder="What is the main point of this document? Or for employee data: 'Show me details about John Smith'",
            height=100
        )
        
        # Store question in session state
        st.session_state.question = question
        
        col1, col2 = st.columns([1, 3])
        with col1:
            submit_button = st.button("Get Answer", use_container_width=True)
        
        if submit_button and question:
            if not groq_api_key:
                st.error("Please enter your Groq API Key in the sidebar")
            else:
                with st.spinner("Generating answer..."):
                    try:
                        llm = ChatGroq(
                            model_name=model_name,
                            temperature=temperature,
                            api_key=groq_api_key
                        )
                        
                        retriever = st.session_state.vector_store.as_retriever(
                            search_kwargs={"k": 4}  # Retrieve top 4 most relevant chunks
                        )
                        
                        qa_chain = RetrievalQA.from_chain_type(
                            llm=llm,
                            chain_type="stuff",
                            retriever=retriever,
                            chain_type_kwargs={"prompt": get_custom_prompt()},
                            return_source_documents=True
                        )
                        
                        result = qa_chain.invoke({"query": question})
                        
                        st.subheader("Answer")
                        st.markdown(result["result"])
                        
                        with st.expander("View Source Documents"):
                            for i, doc in enumerate(result["source_documents"]):
                                st.markdown(f"**Source {i+1}**")
                                st.text(doc.page_content)
                                st.markdown("---")
                    
                    except Exception as e:
                        st.error(f"Error generating answer: {str(e)}")
                        st.info("Make sure your Groq API Key is valid and the service is available")
        # --- Summarization Section ---
        st.subheader("Summarize the Document")
        if st.button("Summarize Document"):
            if not groq_api_key:
                st.error("Please enter your Groq API Key in the sidebar")
            else:
                with st.spinner("Generating summary..."):
                    try:
                        llm = ChatGroq(
                            model_name=model_name,
                            temperature=temperature,
                            api_key=groq_api_key)
                        docs = list(st.session_state.vector_store.docstore._dict.values())
                        summarize_chain = load_summarize_chain(llm, chain_type="map_reduce")
                        summary = summarize_chain.run(docs)

                        st.subheader("Summary")
                        st.success(summary)
                    except Exception as e:
                        st.error(f"Error generating summary: {str(e)}")

    
    elif uploaded_file and not st.session_state.vector_store:
        st.warning("Document processing failed. Please check if the document contains readable text.")

if __name__ == "__main__":
    main()